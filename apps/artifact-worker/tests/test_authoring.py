from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook, load_workbook
from pydantic import ValidationError

from src.app import author
from src.authoring import author_artifacts
from src.contracts import ArtifactWorkerRequest, UploadedArtifact
from src.document_author import create_document, edit_document
from src.presentation_author import create_presentation, edit_presentation
from src.process import convert_with_libreoffice
from src.spreadsheet_author import create_spreadsheet, edit_spreadsheet
from src.validation import validate_artifact

DOCUMENT_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class ArtifactAuthoringTest(unittest.TestCase):
    def test_same_output_filename_never_overwrites_the_downloaded_source(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source_docx = root / "source.docx"
            create_document(
                {
                    "title": "Source",
                    "blocks": [{"type": "paragraph", "text": "Original page"}],
                    "orientation": "portrait",
                    "pageSize": "a4",
                },
                source_docx,
            )
            source_pdf = convert_with_libreoffice(
                source_docx, root / "source-pdf", "pdf"
            )
            request = ArtifactWorkerRequest.model_validate(
                {
                    "callbackToken": "c" * 32,
                    "callbackUrl": "https://example.com/callback",
                    "jobId": "job-same-filename",
                    "operation": {
                        "kind": "pdf_edit",
                        "filename": "report.pdf",
                        "source": {
                            "filename": "report.pdf",
                            "mediaType": "application/pdf",
                            "storageId": "storage-source",
                        },
                        "edits": [{"type": "reorder_pages", "pageNumbers": [1]}],
                    },
                    "sources": [
                        {
                            "downloadUrl": "https://example.com/source",
                            "filename": "report.pdf",
                            "mediaType": "application/pdf",
                            "storageId": "storage-source",
                        }
                    ],
                    "uploads": [
                        {
                            "filename": "report.pdf",
                            "format": "pdf",
                            "mediaType": "application/pdf",
                            "uploadUrl": "https://example.com/upload",
                        }
                    ],
                }
            )
            working_directory = root / "working"
            working_directory.mkdir()
            with patch(
                "src.authoring.download_source",
                side_effect=lambda _source, destination: destination.write_bytes(
                    source_pdf.read_bytes()
                ),
            ):
                outputs = author_artifacts(request, working_directory)

            self.assertEqual(outputs["report.pdf"], working_directory / "report.pdf")
            self.assertTrue(outputs["report.pdf"].is_file())

    def test_set_title_does_not_replace_an_ordinary_first_paragraph(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            edited = root / "edited.docx"
            document = Document()
            document.add_paragraph("Keep this body paragraph.")
            document.save(source)

            edit_document(
                source,
                [{"type": "set_title", "title": "Updated metadata title"}],
                edited,
            )

            result = Document(edited)
            self.assertEqual(result.core_properties.title, "Updated metadata title")
            self.assertEqual(result.paragraphs[0].text, "Keep this body paragraph.")

    def test_worker_reports_partial_uploads_on_failure(self) -> None:
        request = ArtifactWorkerRequest.model_validate(
            {
                "callbackToken": "c" * 32,
                "callbackUrl": "https://example.com/callback",
                "jobId": "job-1",
                "operation": {
                    "kind": "document_create",
                    "document": {
                        "title": "Report",
                        "blocks": [{"type": "paragraph", "text": "Body"}],
                        "orientation": "portrait",
                        "pageSize": "a4",
                    },
                    "outputs": [
                        {"filename": "report.docx", "format": "docx"},
                        {"filename": "report.pdf", "format": "pdf"},
                    ],
                },
                "sources": [],
                "uploads": [
                    {
                        "filename": "report.docx",
                        "format": "docx",
                        "mediaType": DOCUMENT_MEDIA_TYPE,
                        "uploadUrl": "https://example.com/upload-docx",
                    },
                    {
                        "filename": "report.pdf",
                        "format": "pdf",
                        "mediaType": "application/pdf",
                        "uploadUrl": "https://example.com/upload-pdf",
                    },
                ],
            }
        )
        partial_output = UploadedArtifact(
            filename="report.docx",
            media_type=DOCUMENT_MEDIA_TYPE,
            sha256="a" * 64,
            size_bytes=12,
            storage_id="storage-1",
        )
        with (
            patch.dict("os.environ", {"ARTIFACT_WORKER_SECRET": "s" * 32}),
            patch(
                "src.app.author_artifacts",
                return_value={
                    "report.docx": Path("report.docx"),
                    "report.pdf": Path("report.pdf"),
                },
            ),
            patch(
                "src.app.upload_artifact",
                side_effect=[partial_output, ValueError("second upload failed")],
            ),
            patch("src.app.fail_job") as fail_job,
        ):
            with self.assertRaisesRegex(ValueError, "second upload failed"):
                author(request, f"Bearer {'s' * 32}")

        fail_job.assert_called_once_with(
            "https://example.com/callback",
            "c" * 32,
            "Artifact authoring failed: second upload failed",
            "job-1",
            [partial_output],
            "s" * 32,
        )

    def test_worker_boundary_rejects_unstructured_operations(self) -> None:
        with self.assertRaises(ValidationError):
            ArtifactWorkerRequest.model_validate(
                {
                    "callbackToken": "x" * 32,
                    "callbackUrl": "https://example.com/callback",
                    "jobId": "job-1",
                    "operation": {"kind": "document_create", "document": {}},
                    "sources": [],
                    "uploads": [
                        {
                            "filename": "report.docx",
                            "format": "docx",
                            "mediaType": "application/octet-stream",
                            "uploadUrl": "https://example.com/upload",
                        }
                    ],
                }
            )

    def test_validation_rejects_excessive_workbook_sheet_count(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            workbook_path = Path(directory) / "oversized.xlsx"
            workbook = Workbook()
            for index in range(20):
                workbook.create_sheet(f"Sheet {index + 2}")
            workbook.save(workbook_path)

            with self.assertRaisesRegex(ValueError, "20-sheet validation limit"):
                validate_artifact(workbook_path, "xlsx")

    def test_document_create_edit_pdf_and_visual_validation(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            created = root / "report.docx"
            edited = root / "report-edited.docx"
            create_document(
                {
                    "title": "Quarterly report",
                    "subtitle": "Validated output",
                    "blocks": [
                        {"type": "heading", "level": 1, "text": "Summary"},
                        {"type": "paragraph", "text": "Revenue was 1200."},
                        {
                            "type": "table",
                            "headers": ["Month", "Revenue"],
                            "rows": [["January", "1200"], ["February", "1450"]],
                        },
                    ],
                    "orientation": "portrait",
                    "pageSize": "a4",
                },
                created,
            )
            edit_document(
                created,
                [
                    {
                        "type": "replace_text",
                        "find": "1200",
                        "replace": "1,200",
                        "replaceAll": True,
                    },
                    {
                        "type": "append_blocks",
                        "blocks": [
                            {"type": "paragraph", "text": "Approved for release."}
                        ],
                    },
                ],
                edited,
            )
            validate_artifact(edited, "docx")
            pdf = convert_with_libreoffice(edited, root / "pdf", "pdf")
            validate_artifact(pdf, "pdf")

    def test_spreadsheet_create_edit_chart_and_visual_validation(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            created = root / "metrics.xlsx"
            edited = root / "metrics-edited.xlsx"
            create_spreadsheet(
                [
                    {
                        "name": "Revenue",
                        "rows": [
                            ["Month", "Revenue"],
                            ["January", 1200],
                            ["February", 1450],
                        ],
                        "frozenRows": 1,
                        "charts": [
                            {
                                "type": "bar",
                                "title": "Monthly revenue",
                                "categoryColumn": "Month",
                                "dataColumns": ["Revenue"],
                                "position": "D2",
                            }
                        ],
                    }
                ],
                created,
            )
            edit_spreadsheet(
                created,
                [
                    {
                        "type": "append_rows",
                        "sheet": "Revenue",
                        "rows": [["March", 1600]],
                    },
                    {
                        "type": "set_cell",
                        "sheet": "Revenue",
                        "cell": "B5",
                        "value": "=SUM(B2:B4)",
                    },
                ],
                edited,
            )
            validate_artifact(edited, "xlsx")
            workbook = load_workbook(edited)
            chart_series = workbook["Revenue"]._charts[0].series[0]
            self.assertEqual(chart_series.cat.numRef.f, "'Revenue'!$A$2:$A$4")
            self.assertEqual(chart_series.val.numRef.f, "'Revenue'!$B$2:$B$4")

    def test_presentation_create_insert_delete_and_visual_validation(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            created = root / "briefing.pptx"
            edited = root / "briefing-edited.pptx"
            title_slide = {
                "layout": "title",
                "title": "Launch briefing",
                "subtitle": "Validated presentation",
            }
            content_slide = {
                "layout": "content",
                "title": "Milestones",
                "bullets": ["Private beta", "Public launch"],
            }
            create_presentation(
                {"title": "Launch briefing", "slides": [title_slide, content_slide]},
                created,
            )
            edit_presentation(
                created,
                [
                    {
                        "type": "insert_slides",
                        "afterSlide": 1,
                        "slides": [
                            {
                                "layout": "section",
                                "title": "Product",
                                "subtitle": "What ships first",
                            }
                        ],
                    },
                    {"type": "delete_slides", "slideNumbers": [3]},
                ],
                edited,
            )
            validate_artifact(edited, "pptx")


if __name__ == "__main__":
    unittest.main()
