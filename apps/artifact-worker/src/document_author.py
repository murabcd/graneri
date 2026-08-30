from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

from .text_runs import replace_text_in_runs


def _configure_document(document: DocumentType, spec: dict[str, Any]) -> None:
    section = document.sections[0]
    if spec.get("pageSize") == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    if spec.get("orientation") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (
            section.page_height,
            section.page_width,
        )
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Liberation Sans"
    normal_style.font.size = Pt(10.5)
    for level in range(1, 4):
        heading = document.styles[f"Heading {level}"]
        heading.font.name = "Liberation Sans"
        heading.font.color.rgb = None


def _add_table(document: DocumentType, block: dict[str, Any]) -> None:
    headers = block["headers"]
    rows = block["rows"]
    if any(len(row) != len(headers) for row in rows):
        raise ValueError("Document table rows must match the header width.")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        run = paragraph.add_run(str(header))
        run.bold = True
    for source_row in rows:
        row = table.add_row()
        for index in range(len(headers)):
            row.cells[index].text = (
                str(source_row[index]) if index < len(source_row) else ""
            )


def add_document_blocks(document: DocumentType, blocks: list[dict[str, Any]]) -> None:
    for block in blocks:
        block_type = block["type"]
        if block_type == "heading":
            document.add_heading(block["text"], level=block["level"])
        elif block_type == "paragraph":
            document.add_paragraph(block["text"])
        elif block_type == "bullet_list":
            for item in block["items"]:
                document.add_paragraph(item, style="List Bullet")
        elif block_type == "numbered_list":
            for item in block["items"]:
                document.add_paragraph(item, style="List Number")
        elif block_type == "table":
            _add_table(document, block)
        elif block_type == "page_break":
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        else:
            raise ValueError(f"Unsupported document block: {block_type}")


def create_document(spec: dict[str, Any], destination: Path) -> None:
    document = Document()
    _configure_document(document, spec)
    document.core_properties.title = spec["title"]
    if spec.get("author"):
        document.core_properties.author = spec["author"]
    document.add_heading(spec["title"], level=0)
    if spec.get("subtitle"):
        subtitle = document.add_paragraph(spec["subtitle"])
        subtitle.style = document.styles["Subtitle"]
    add_document_blocks(document, spec["blocks"])
    document.save(destination)


def _iter_paragraphs(document: DocumentType):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _replace_text(
    document: DocumentType, find: str, replace: str, replace_all: bool
) -> int:
    replacement_count = 0
    for paragraph in _iter_paragraphs(document):
        replacements = replace_text_in_runs(
            paragraph.runs,
            find,
            replace,
            replace_all,
        )
        replacement_count += replacements
        if replacements > 0 and not replace_all:
            return replacement_count
    return replacement_count


def edit_document(source: Path, edits: list[dict[str, Any]], destination: Path) -> None:
    document = Document(source)
    for edit in edits:
        edit_type = edit["type"]
        if edit_type == "append_blocks":
            add_document_blocks(document, edit["blocks"])
        elif edit_type == "replace_text":
            replacements = _replace_text(
                document,
                edit["find"],
                edit["replace"],
                edit.get("replaceAll", False),
            )
            if replacements == 0:
                raise ValueError(f"Document text was not found: {edit['find']}")
        elif edit_type == "set_title":
            document.core_properties.title = edit["title"]
            for paragraph in document.paragraphs:
                if paragraph.style.name == "Title":
                    paragraph.text = edit["title"]
                    break
        else:
            raise ValueError(f"Unsupported document edit: {edit_type}")
    document.save(destination)
